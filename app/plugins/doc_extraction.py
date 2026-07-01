"""
Pure-function text extraction pipeline for 6 file formats.

Provides format-agnostic text extraction with encoding detection,
SHA-256 fingerprinting, and scanned-PDF flagging. No plugin
dependencies — every function is independently testable.

Exports:
    EXTRACTORS      — dispatch table mapping extension → extractor callable
    extract_text    — main entry point, returns {text, encoding, needs_ocr, sha256, error}
    compute_sha256  — SHA-256 hex digest helper
    detect_encoding — charset-normalizer wrapper
    extract_docx / extract_pdf / extract_doc / extract_rst / extract_drawio / extract_graphml
"""

from __future__ import annotations

import hashlib
import logging
import os
import xml.etree.ElementTree as ET
from typing import Callable

logger = logging.getLogger("doc_extraction")


# ═══════════════════════════════════════════════════════════════
#  SHA-256 helper
# ═══════════════════════════════════════════════════════════════

def compute_sha256(file_path: str) -> str:
    """Return lowercase hex SHA-256 digest of *file_path*.

    Reads in 64 KB chunks to stay within memory budget for large files.
    Returns ``""`` on ``FileNotFoundError`` (logged at warning level).
    """
    sha = hashlib.sha256()
    try:
        with open(file_path, "rb") as fh:
            while True:
                chunk = fh.read(65536)  # 64 KB
                if not chunk:
                    break
                sha.update(chunk)
    except FileNotFoundError:
        logger.warning("SHA-256: file not found: %s", file_path)
        return ""
    except OSError as exc:
        logger.warning("SHA-256: OS error reading %s: %s", file_path, exc)
        return ""
    return sha.hexdigest()


# ═══════════════════════════════════════════════════════════════
#  Encoding detection
# ═══════════════════════════════════════════════════════════════

def detect_encoding(text: str) -> str:
    """Detect the character encoding of *text* via charset-normalizer.

    Encodes *text* to UTF-8 bytes (with ``replace`` for safety), then
    passes through ``charset_normalizer.from_bytes().best()``.  Falls
    back to ``"utf_8"`` when detection fails.

    Returns the encoding name as a string, e.g. ``"utf_8"``, ``"windows_1257"``.
    """
    try:
        from charset_normalizer import from_bytes  # lazy import — pure function
    except ImportError:
        logger.warning("charset-normalizer not installed; defaulting to utf_8")
        return "utf_8"

    try:
        raw = text.encode("utf-8", errors="replace")
        best = from_bytes(raw).best()
        if best is not None:
            encoding = best.encoding or "utf_8"
        else:
            encoding = "utf_8"
        logger.debug("detect_encoding → %s", encoding)
        return encoding
    except Exception as exc:
        logger.warning("Encoding detection failed: %s; defaulting to utf_8", exc)
        return "utf_8"


# ═══════════════════════════════════════════════════════════════
#  Format extractors — signature: (file_path: str) -> str
# ═══════════════════════════════════════════════════════════════

def extract_docx(file_path: str) -> str:
    """Extract text from a .docx file via python-docx.

    Iterates paragraphs and tables; each table row is joined with tabs.
    """
    try:
        from docx import Document  # lazy import
    except ImportError:
        logger.warning("python-docx not installed; cannot extract %s", file_path)
        return ""

    try:
        doc = Document(file_path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append("\t".join(cells))

        return "\n".join(parts)
    except Exception as exc:
        logger.warning("Failed to extract .docx: %s: %s", file_path, exc)
        return ""


def extract_pdf(file_path: str) -> str:
    """Extract text from a .pdf file.

    Primary path:  ``pdfplumber``.
    Fallback path: ``pypdf.PdfReader`` (used when pdfplumber fails).

    If the PDF has pages but extracted text is shorter than 20 characters
    the function returns ``""`` so the caller can set ``needs_ocr: true``.
    """
    # ── Primary: pdfplumber ─────────────────────────────────
    try:
        import pdfplumber  # lazy import

        with pdfplumber.open(file_path) as pdf:
            pages_text: list[str] = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages_text.append(page_text)
            text = "\n".join(pages_text).strip()
            page_count = len(pdf.pages)

        # Scanned-PDF detection: pages exist but no meaningful text
        if page_count > 0 and len(text) < 20:
            logger.debug("PDF appears to be scanned (pages=%d, text_len=%d): %s",
                         page_count, len(text), file_path)
            return ""

        return text
    except Exception:
        pass  # fall through to pypdf fallback

    # ── Fallback: pypdf ─────────────────────────────────────
    try:
        from pypdf import PdfReader  # lazy import

        reader = PdfReader(file_path)
        pages_text: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages_text.append(page_text)
        text = "\n".join(pages_text).strip()
        page_count = len(reader.pages)

        if page_count > 0 and len(text) < 20:
            logger.debug("PDF appears to be scanned (pages=%d, text_len=%d): %s",
                         page_count, len(text), file_path)
            return ""

        return text
    except Exception as exc:
        logger.warning("Failed to extract .pdf: %s: %s", file_path, exc)
        return ""


def extract_doc(file_path: str) -> str:
    """Extract text from a legacy .doc file via doc2txt."""
    try:
        import doc2txt  # lazy import
    except ImportError:
        logger.warning("doc2txt not installed (antiword may be missing); cannot extract %s", file_path)
        return ""

    try:
        return doc2txt.extract_text(file_path)
    except Exception as exc:
        logger.warning("Failed to extract .doc: %s: %s", file_path, exc)
        return ""


def extract_rst(file_path: str) -> str:
    """Extract text from a reStructuredText file via docutils.

    Parses into a document tree, then recursively walks all nodes
    extracting ``node.astext()`` from text-bearing nodes.
    """
    try:
        import docutils.core  # lazy import
        import docutils.io
        import docutils.nodes
    except ImportError:
        logger.warning("docutils not installed; cannot extract %s", file_path)
        return ""

    try:
        with open(file_path, encoding="utf-8", errors="replace") as fh:
            source = fh.read()

        pub = docutils.core.Publisher(
            source_class=docutils.io.StringInput,
            destination_class=docutils.io.NullOutput,
        )
        pub.set_components("standalone", "restructuredtext", "null")
        pub.process_programmatic_settings(None, None, None)
        pub.set_source(source=source)
        pub.publish()
        document = pub.document

        texts: list[str] = []

        def _walk(node: docutils.nodes.Node) -> None:
            if isinstance(node, (docutils.nodes.Text,)):
                t = node.astext()
                if t.strip():
                    texts.append(t)
            for child in node.children:
                _walk(child)

        _walk(document)
        return "\n".join(texts)
    except Exception as exc:
        logger.warning("Failed to extract .rst: %s: %s", file_path, exc)
        return ""


def extract_drawio(file_path: str) -> str:
    """Extract text from a .drawio (draw.io) XML file.

    Collects ``value`` attributes from every ``mxCell`` element.
    """
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()
        texts: list[str] = []

        for cell in root.iter("mxCell"):
            val = cell.get("value", "")
            if val and val.strip():
                texts.append(val.strip())

        return "\n".join(texts)
    except ET.ParseError as exc:
        logger.warning("Failed to extract .drawio (XML parse error): %s: %s", file_path, exc)
        return ""
    except Exception as exc:
        logger.warning("Failed to extract .drawio: %s: %s", file_path, exc)
        return ""


def extract_graphml(file_path: str) -> str:
    """Extract text from a .graphml XML file.

    Looks for ``<data>`` child elements inside ``<node>`` elements using
    the standard GraphML namespace.  Falls back to a no-namespace search
    if the namespace-aware path yields nothing.
    """
    try:
        tree = ET.parse(file_path)
        root = tree.getroot()

        ns = {"graphml": "http://graphml.graphdrawing.org/xmlns"}
        texts: list[str] = []

        # ── Namespace-aware path ────────────────────────────
        for node_elem in root.findall(".//graphml:node", ns):
            for data_elem in node_elem.findall("graphml:data", ns):
                val = (data_elem.text or "").strip()
                if val:
                    texts.append(val)

        # ── No-namespace fallback ───────────────────────────
        if not texts:
            for node_elem in root.findall(".//node"):
                for data_elem in node_elem.findall("data"):
                    val = (data_elem.text or "").strip()
                    if val:
                        texts.append(val)

        return "\n".join(texts)
    except ET.ParseError as exc:
        logger.warning("Failed to extract .graphml (XML parse error): %s: %s", file_path, exc)
        return ""
    except Exception as exc:
        logger.warning("Failed to extract .graphml: %s: %s", file_path, exc)
        return ""


# ═══════════════════════════════════════════════════════════════
#  Format dispatch table
# ═══════════════════════════════════════════════════════════════

EXTRACTORS: dict[str, Callable[[str], str]] = {
    ".docx": extract_docx,
    ".pdf": extract_pdf,
    ".doc": extract_doc,
    ".rst": extract_rst,
    ".drawio": extract_drawio,
    ".graphml": extract_graphml,
}


# ═══════════════════════════════════════════════════════════════
#  Main dispatch function
# ═══════════════════════════════════════════════════════════════

def extract_text(file_path: str) -> dict:
    """Extract text from *file_path*, dispatching by file extension.

    Returns:
        dict with keys:
            text      — extracted text (empty string on failure)
            encoding  — detected encoding name, e.g. ``"utf_8"``
            needs_ocr — ``True`` when PDF has pages > 0 but text < 20 chars
            sha256    — lowercase hex SHA-256 digest of the file
            error     — error message string, or ``None`` on success
    """
    result: dict = {
        "text": "",
        "encoding": "utf_8",
        "needs_ocr": False,
        "sha256": "",
        "error": None,
    }

    # 1. SHA-256 fingerprint
    sha = compute_sha256(file_path)
    result["sha256"] = sha

    # 2. Determine extension
    ext = os.path.splitext(file_path)[1].lower()

    # 3. Lookup extractor
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        result["error"] = f"Unsupported format: {ext}"
        logger.warning("Unsupported format for extraction: %s (%s)", file_path, ext)
        return result

    # 4. Call extractor (must never raise)
    try:
        raw_text = extractor(file_path)
    except Exception as exc:
        logger.warning("Extraction crashed for %s: %s", file_path, exc)
        result["error"] = str(exc)
        return result

    text = raw_text.strip() if raw_text else ""
    result["text"] = text

    # 4b. Detect silent failures — file missing or extractor produced nothing
    if not text and not sha:
        result["error"] = f"File not found or unreadable: {file_path}"
    elif not text and sha and ext != ".pdf":
        # File exists but extraction produced nothing (corrupt / unsupported content)
        result["error"] = f"Extraction produced no text for: {file_path}"

    # 5. Encoding detection
    if text:
        result["encoding"] = detect_encoding(text)
    else:
        result["encoding"] = "utf_8"

    # 6. Scanned-PDF flag
    if ext == ".pdf" and not text:
        # If extract_pdf returned "" and the file has pages with no text,
        # we need a quick page-count check to confirm it's a scanned PDF.
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            if page_count > 0:
                result["needs_ocr"] = True
                logger.debug("Flagged as scanned PDF (pages=%d, text_len=%d): %s",
                             page_count, len(text), file_path)
        except Exception:
            # If we can't even open it with pypdf, it's truly corrupt — not scanned
            if result["error"] is None:
                result["error"] = "PDF extraction failed (possible corrupt file)"

    return result
